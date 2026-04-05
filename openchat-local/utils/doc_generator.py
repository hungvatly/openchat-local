"""
OpenChat Local — Document Generator
Creates Word (.docx), PDF (.pdf), and Excel (.xlsx) files from AI-generated content.
"""
import os
import re
import json
import uuid
from typing import Dict, Optional

OUTPUT_DIR = os.path.join("data", "generated")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def generate_docx(title: str, content: str, filename: str = None) -> Dict:
    """Generate a Word document from text content."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"

        doc.add_heading(title, level=1)

        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("# "):
                doc.add_heading(block[2:], level=1)
            elif block.startswith("## "):
                doc.add_heading(block[3:], level=2)
            elif block.startswith("### "):
                doc.add_heading(block[4:], level=3)
            elif block.startswith("- ") or block.startswith("* "):
                for line in block.split("\n"):
                    line = line.strip()
                    if line.startswith(("- ", "* ")):
                        doc.add_paragraph(line[2:], style="List Bullet")
                    elif line:
                        doc.add_paragraph(line)
            elif re.match(r"^\d+\.\s", block):
                for line in block.split("\n"):
                    line = line.strip()
                    m = re.match(r"^\d+\.\s(.+)", line)
                    if m:
                        doc.add_paragraph(m.group(1), style="List Number")
                    elif line:
                        doc.add_paragraph(line)
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", block)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                doc.add_paragraph(clean)

        fname = filename or f"{uuid.uuid4().hex[:8]}_{_slug(title)}.docx"
        fpath = os.path.join(OUTPUT_DIR, fname)
        doc.save(fpath)
        return {"status": "ok", "path": fpath, "filename": fname, "type": "docx"}
    except ImportError:
        return {"status": "error", "message": "python-docx not installed"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def generate_pdf(title: str, content: str, filename: str = None) -> Dict:
    """Generate a PDF from text content."""
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(6)

        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue

            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", block)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)

            if clean.startswith("# "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(0, 10, clean[2:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)
            elif clean.startswith("## "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 9, clean[3:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            elif clean.startswith("### "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 8, clean[4:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            else:
                pdf.set_font("Helvetica", "", 11)
                pdf.multi_cell(0, 6, clean)
                pdf.ln(3)

        fname = filename or f"{uuid.uuid4().hex[:8]}_{_slug(title)}.pdf"
        fpath = os.path.join(OUTPUT_DIR, fname)
        pdf.output(fpath)
        return {"status": "ok", "path": fpath, "filename": fname, "type": "pdf"}
    except ImportError:
        return {"status": "error", "message": "fpdf2 not installed. Run: pip install fpdf2"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def generate_xlsx(title: str, content: str, filename: str = None) -> Dict:
    """Generate an Excel file from tabular content."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]

        header_font = Font(bold=True, size=12)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=11, color="FFFFFF")

        lines = [l.strip() for l in content.strip().split("\n") if l.strip()]

        # Try to detect table format (pipe-separated or CSV-like)
        if any("|" in l for l in lines):
            # Markdown table
            row_num = 1
            for line in lines:
                if set(line.replace("|", "").strip()) <= set("- :"):
                    continue  # skip separator row
                cells = [c.strip() for c in line.strip("|").split("|")]
                for col, val in enumerate(cells, 1):
                    cell = ws.cell(row=row_num, column=col, value=val)
                    if row_num == 1:
                        cell.font = header_font_white
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal="center")
                row_num += 1
        elif any("," in l for l in lines):
            # CSV-like
            import csv
            import io
            reader = csv.reader(io.StringIO(content))
            for row_num, row in enumerate(reader, 1):
                for col, val in enumerate(row, 1):
                    cell = ws.cell(row=row_num, column=col, value=val.strip())
                    if row_num == 1:
                        cell.font = header_font_white
                        cell.fill = header_fill
        else:
            # Just dump lines into column A
            ws.cell(row=1, column=1, value=title).font = header_font
            for i, line in enumerate(lines, 2):
                ws.cell(row=i, column=1, value=line)

        # Auto-width columns
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

        fname = filename or f"{uuid.uuid4().hex[:8]}_{_slug(title)}.xlsx"
        fpath = os.path.join(OUTPUT_DIR, fname)
        wb.save(fpath)
        return {"status": "ok", "path": fpath, "filename": fname, "type": "xlsx"}
    except ImportError:
        return {"status": "error", "message": "openpyxl not installed. Run: pip install openpyxl"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def detect_and_generate(ai_response: str, user_message: str) -> Optional[Dict]:
    """Check if AI response contains content meant for a file and generate it."""
    msg_lower = user_message.lower()

    # Detect document creation intent from the user's message
    if any(kw in msg_lower for kw in ["create a word", "write a docx", "make a word doc", "generate a document", "create a report"]):
        title = _extract_title(user_message, ai_response)
        return generate_docx(title, ai_response)

    if any(kw in msg_lower for kw in ["create a pdf", "make a pdf", "generate a pdf", "export as pdf", "save as pdf"]):
        title = _extract_title(user_message, ai_response)
        return generate_pdf(title, ai_response)

    if any(kw in msg_lower for kw in ["create a spreadsheet", "make an excel", "create an xlsx", "generate a table", "create a csv", "make a spreadsheet"]):
        title = _extract_title(user_message, ai_response)
        return generate_xlsx(title, ai_response)

    return None


def _slug(text: str) -> str:
    slug = re.sub(r"[^\w\s-]", "", text.lower())
    return re.sub(r"[\s-]+", "_", slug).strip("_")[:40]


def _extract_title(user_msg: str, ai_response: str) -> str:
    """Try to extract a sensible title."""
    for prefix in ["create a ", "make a ", "generate a ", "write a "]:
        if prefix in user_msg.lower():
            after = user_msg.lower().split(prefix, 1)[1]
            # remove format words
            for fmt in ["word doc", "docx", "pdf", "spreadsheet", "excel", "xlsx", "document", "report", "table"]:
                after = after.replace(fmt, "").strip()
            if after and len(after) > 3:
                return after.strip(" .,!?")[:60].title()
    # Fallback: first line of AI response
    first_line = ai_response.strip().split("\n")[0]
    clean = re.sub(r"^#+\s*", "", first_line)
    return clean[:60] if clean else "Document"
