"""
OpenChat Local — Template Document Generator
Upload a form/template document, AI analyzes its structure,
then generates a new document with the same layout but filled with your data.
"""
import os
import re
import json
import uuid
import shutil
from typing import Dict, List, Optional, Tuple

TEMPLATE_DIR = os.path.join("data", "templates")
OUTPUT_DIR = os.path.join("data", "generated")
os.makedirs(TEMPLATE_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ── Template Storage ───────────────────────────────────

def save_template(filepath: str, original_name: str) -> Dict:
    """Save an uploaded file as a reusable template."""
    template_id = uuid.uuid4().hex[:8]
    ext = os.path.splitext(original_name)[1].lower()
    safe_name = re.sub(r"[^\w\-.]", "_", original_name)
    dest = os.path.join(TEMPLATE_DIR, f"{template_id}_{safe_name}")
    shutil.copy2(filepath, dest)

    # Extract structure for preview
    structure = extract_template_structure(dest)

    # Save metadata
    meta = {
        "id": template_id,
        "name": original_name,
        "path": dest,
        "extension": ext,
        "structure": structure,
    }
    meta_path = dest + ".json"
    with open(meta_path, "w") as f:
        json.dump(meta, f, indent=2)

    return meta


def list_templates() -> List[Dict]:
    """List all saved templates."""
    templates = []
    for fname in sorted(os.listdir(TEMPLATE_DIR)):
        if fname.endswith(".json"):
            try:
                with open(os.path.join(TEMPLATE_DIR, fname), "r") as f:
                    meta = json.load(f)
                templates.append({
                    "id": meta["id"],
                    "name": meta["name"],
                    "extension": meta["extension"],
                    "fields": len(meta.get("structure", {}).get("fields", [])),
                })
            except Exception:
                continue
    return templates


def get_template(template_id: str) -> Optional[Dict]:
    """Get a template by ID."""
    for fname in os.listdir(TEMPLATE_DIR):
        if fname.endswith(".json") and fname.startswith(template_id):
            with open(os.path.join(TEMPLATE_DIR, fname), "r") as f:
                return json.load(f)
    return None


def delete_template(template_id: str) -> bool:
    """Delete a template."""
    for fname in os.listdir(TEMPLATE_DIR):
        if fname.startswith(template_id):
            os.remove(os.path.join(TEMPLATE_DIR, fname))
    return True


# ── Structure Extraction ───────────────────────────────

def extract_template_structure(filepath: str) -> Dict:
    """Extract the structure/layout of a template document."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".docx":
        return _extract_docx_structure(filepath)
    elif ext == ".pdf":
        return _extract_pdf_structure(filepath)
    elif ext in (".txt", ".md"):
        return _extract_text_structure(filepath)
    else:
        return {"format": ext, "fields": [], "raw_text": ""}


def _extract_docx_structure(filepath: str) -> Dict:
    """Extract structure from a Word document."""
    try:
        from docx import Document

        doc = Document(filepath)
        fields = []
        sections = []
        full_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text.append(text)

            # Detect headings
            if para.style.name.startswith("Heading"):
                sections.append({"type": "heading", "level": para.style.name, "text": text})

            # Detect fields: lines with colons, underscores, brackets, or blank spaces
            if ":" in text:
                parts = text.split(":", 1)
                label = parts[0].strip()
                value = parts[1].strip() if len(parts) > 1 else ""
                if label and len(label) < 60:
                    fields.append({"label": label, "current_value": value, "type": "text"})
            elif "____" in text or "......" in text or "[" in text:
                clean = re.sub(r"[_.\[\]]", "", text).strip()
                if clean:
                    fields.append({"label": clean, "current_value": "", "type": "text"})

        # Detect tables
        tables_info = []
        for i, table in enumerate(doc.tables):
            headers = []
            rows_data = []
            for j, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if j == 0:
                    headers = cells
                else:
                    rows_data.append(cells)
            tables_info.append({"headers": headers, "rows": rows_data, "index": i})

        return {
            "format": "docx",
            "fields": fields,
            "sections": sections,
            "tables": tables_info,
            "raw_text": "\n".join(full_text),
        }
    except Exception as e:
        return {"format": "docx", "fields": [], "error": str(e), "raw_text": ""}


def _extract_pdf_structure(filepath: str) -> Dict:
    """Extract structure from a PDF."""
    try:
        # Try pymupdf
        try:
            import fitz
            doc = fitz.open(filepath)
            full_text = ""
            for page in doc:
                full_text += page.get_text() + "\n"
            doc.close()

            # If no text, try OCR
            if not full_text.strip():
                try:
                    import pytesseract
                    from PIL import Image
                    import io
                    doc = fitz.open(filepath)
                    for page in doc:
                        pix = page.get_pixmap(dpi=150)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        full_text += pytesseract.image_to_string(img) + "\n"
                    doc.close()
                except ImportError:
                    pass
        except ImportError:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            full_text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    full_text += t + "\n"

        if not full_text.strip():
            return {"format": "pdf", "fields": [], "raw_text": "", "error": "Could not extract text"}

        # Detect fields from text
        fields = _detect_fields_from_text(full_text)

        return {
            "format": "pdf",
            "fields": fields,
            "raw_text": full_text,
        }
    except Exception as e:
        return {"format": "pdf", "fields": [], "error": str(e), "raw_text": ""}


def _extract_text_structure(filepath: str) -> Dict:
    """Extract structure from plain text/markdown."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    fields = _detect_fields_from_text(text)
    return {"format": "txt", "fields": fields, "raw_text": text}


def _detect_fields_from_text(text: str) -> List[Dict]:
    """Detect fillable fields from raw text."""
    fields = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Pattern: "Label: value" or "Label: ___"
        if ":" in line:
            parts = line.split(":", 1)
            label = parts[0].strip()
            value = parts[1].strip()
            if label and len(label) < 60 and not label.startswith("http"):
                # Clean placeholder values
                clean_val = re.sub(r"[_.\[\]]+", "", value).strip()
                fields.append({"label": label, "current_value": clean_val, "type": "text"})
        # Pattern: lines with blanks
        elif "____" in line or "......" in line:
            clean = re.sub(r"[_.\[\]]", "", line).strip()
            if clean:
                fields.append({"label": clean, "current_value": "", "type": "text"})
    return fields


# ── Build Prompt for AI ────────────────────────────────

def build_fill_prompt(template: Dict, user_instructions: str) -> str:
    """Build a prompt that tells the AI to fill in the template."""
    structure = template.get("structure", {})
    fields = structure.get("fields", [])
    raw_text = structure.get("raw_text", "")
    tables = structure.get("tables", [])

    prompt_parts = [
        "You are filling in a document template. Below is the original template structure.",
        "Your task is to generate a COMPLETE document that follows the EXACT same layout, structure, headings, and formatting as the template, but with the fields filled in using the information provided by the user.",
        "",
        "IMPORTANT RULES:",
        "- Keep the EXACT same document structure, section order, and formatting",
        "- Fill in all blank fields and placeholders with appropriate information",
        "- If information is not provided by the user, use reasonable placeholder text like '[To be determined]'",
        "- Preserve any tables in the same format",
        "- Output the filled document as clean text with markdown formatting (headings, bold, tables, lists)",
        "",
    ]

    if fields:
        prompt_parts.append("## Detected Fields in Template:")
        for f in fields:
            val_note = f" (currently: {f['current_value']})" if f['current_value'] else " (blank)"
            prompt_parts.append(f"- {f['label']}{val_note}")
        prompt_parts.append("")

    if tables:
        prompt_parts.append("## Tables in Template:")
        for t in tables:
            prompt_parts.append(f"Table with headers: {', '.join(t['headers'])}")
            prompt_parts.append(f"  Existing rows: {len(t['rows'])}")
        prompt_parts.append("")

    prompt_parts.append("## Full Template Text:")
    # Truncate if very long
    template_text = raw_text[:4000] if len(raw_text) > 4000 else raw_text
    prompt_parts.append(template_text)
    prompt_parts.append("")
    prompt_parts.append("## User's Instructions:")
    prompt_parts.append(user_instructions)
    prompt_parts.append("")
    prompt_parts.append("Now generate the complete filled document following the template structure exactly. Use markdown formatting.")

    return "\n".join(prompt_parts)


# ── Generate Filled Document ───────────────────────────

def generate_from_template(template_id: str, ai_content: str, output_format: str = None) -> Dict:
    """Generate a filled document from AI content, matching template format."""
    template = get_template(template_id)
    if not template:
        return {"status": "error", "message": "Template not found"}

    ext = output_format or template.get("extension", ".docx")
    title = re.sub(r"\.[^.]+$", "", template["name"])
    fname = f"{uuid.uuid4().hex[:8]}_filled_{re.sub(r'[^a-zA-Z0-9]', '_', title)}{ext}"

    if ext == ".docx":
        return _generate_filled_docx(template, ai_content, fname)
    elif ext == ".pdf":
        from utils.doc_generator import generate_pdf
        return generate_pdf(f"Filled: {title}", ai_content, fname)
    else:
        # Fallback: save as text
        fpath = os.path.join(OUTPUT_DIR, fname.replace(ext, ".md"))
        with open(fpath, "w", encoding="utf-8") as f:
            f.write(ai_content)
        return {"status": "ok", "path": fpath, "filename": os.path.basename(fpath), "type": "md"}


def _generate_filled_docx(template: Dict, ai_content: str, filename: str) -> Dict:
    """Generate a filled Word document preserving template styling where possible."""
    try:
        from docx import Document
        from docx.shared import Pt

        # Try to use original template as base for styling
        template_path = template.get("path", "")
        if os.path.exists(template_path) and template_path.endswith(".docx"):
            try:
                doc = Document(template_path)
                # Clear all paragraphs but keep styles
                for para in doc.paragraphs:
                    para.clear()
                # Clear tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = ""
            except Exception:
                doc = Document()
        else:
            doc = Document()

        # Parse AI content and build document
        style = doc.styles["Normal"]
        style.font.size = Pt(11)

        # Remove empty paragraphs from cleared template
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            p = doc.paragraphs[-1]._element
            p.getparent().remove(p)

        for block in ai_content.split("\n\n"):
            block = block.strip()
            if not block:
                continue

            if block.startswith("# "):
                doc.add_heading(block[2:], level=1)
            elif block.startswith("## "):
                doc.add_heading(block[3:], level=2)
            elif block.startswith("### "):
                doc.add_heading(block[4:], level=3)
            elif "|" in block and block.count("|") >= 2:
                # Markdown table
                lines = [l.strip() for l in block.split("\n") if l.strip()]
                table_rows = []
                for line in lines:
                    if set(line.replace("|", "").strip()) <= set("- :"):
                        continue
                    cells = [c.strip() for c in line.strip("|").split("|")]
                    table_rows.append(cells)
                if table_rows:
                    cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=cols)
                    table.style = "Table Grid"
                    for i, row_data in enumerate(table_rows):
                        for j, val in enumerate(row_data):
                            if j < cols:
                                table.rows[i].cells[j].text = val
            elif block.startswith("- ") or block.startswith("* "):
                for line in block.split("\n"):
                    line = line.strip()
                    if line.startswith(("- ", "* ")):
                        doc.add_paragraph(line[2:], style="List Bullet")
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", block)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                doc.add_paragraph(clean)

        fpath = os.path.join(OUTPUT_DIR, filename)
        doc.save(fpath)
        return {"status": "ok", "path": fpath, "filename": filename, "type": "docx"}
    except Exception as e:
        return {"status": "error", "message": str(e)}
